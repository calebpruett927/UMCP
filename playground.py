"""
RCFT Playground: an interactive what-if tool for tuning UMCP parameters.

This script reads a CSV file and computes invariants under different parameter settings.
It can generate a density plot of omega vs curvature and export a docx report summarizing the run.
"""
import argparse
import csv
import json
import os
from typing import List
import matplotlib.pyplot as plt
from statistics import mean
from docx import Document

# Reuse the core computation from Live Gauge
from live_gauge.core import compute_invariants

def parse_csv(file_path: str, channel: str) -> List[float]:
    x_raw = []
    with open(file_path, 'r', newline='') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            try:
                value = float(row[channel])
                x_raw.append(value)
            except (KeyError, ValueError):
                continue
    return x_raw

def density_plot(results, out_path: str):
    omegas = [r['omega'] for r in results]
    Cs = [r['C'] for r in results]
    plt.figure(figsize=(6, 4))
    plt.hexbin(omegas, Cs, gridsize=50, cmap='Blues', mincnt=1)
    plt.xlabel('omega')
    plt.ylabel('C')
    plt.title('Density of omega vs curvature (C)')
    plt.colorbar(label='Counts')
    # Gate lines
    plt.axvline(0.038, color='orange', linestyle='--', label='Watch gate')
    plt.axvline(0.30, color='red', linestyle='--', label='Collapse gate')
    plt.axhline(0.14, color='green', linestyle='--', label='C gate')
    plt.legend()
    plt.tight_layout()
    plt.savefig(out_path)
    plt.close()

def export_docx(summary: dict, plot_path: str, out_docx: str):
    doc = Document()
    doc.add_heading('RCFT Playground Report', level=1)
    doc.add_paragraph(f"Samples processed: {summary['n_samples']}")
    doc.add_paragraph(f"Parameters: epsilon={summary['epsilon']}, K={summary['K']}, a={summary['a']}, b={summary['b']}")
    doc.add_paragraph('Regime counts:')
    for regime, count in summary['regime_counts'].items():
        doc.add_paragraph(f"- {regime}: {count}", style='List Bullet')
    doc.add_paragraph(f"Mean omega: {summary['mean_omega']:.4f}")
    doc.add_paragraph(f"Mean curvature C: {summary['mean_C']:.4f}")
    doc.add_paragraph(f"Mean kappa: {summary['mean_kappa']:.4f}")

    doc.add_paragraph('See the following density plot for omega vs C:')
    # Add image
    doc.add_picture(plot_path, width=doc.sections[0].page_width * 0.7)
    doc.save(out_docx)

def main():
    parser = argparse.ArgumentParser(description='RCFT Playground')
    parser.add_argument('--csv', required=True, help='CSV file with data')
    parser.add_argument('--channel', required=True, help='Column name for the data')
    parser.add_argument('--a', type=float, required=True, help='Frozen intercept (a)')
    parser.add_argument('--b', type=float, required=True, help='Frozen scale (b)')
    parser.add_argument('--epsilon', type=float, default=1e-8, help='Small epsilon for entropy calculation')
    parser.add_argument('--K', type=int, default=3, help='Curvature window length')
    parser.add_argument('--out_dir', default='playground_out', help='Output directory')
    args = parser.parse_args()

    os.makedirs(args.out_dir, exist_ok=True)
    x_raw = parse_csv(args.csv, args.channel)
    results = compute_invariants(x_raw, a=args.a, b=args.b, epsilon=args.epsilon, K=args.K)

    regime_counts = {}
    for r in results:
        regime_counts[r['regime']] = regime_counts.get(r['regime'], 0) + 1
    summary = {
        'n_samples': len(results),
        'epsilon': args.epsilon,
        'K': args.K,
        'a': args.a,
        'b': args.b,
        'regime_counts': regime_counts,
        'mean_omega': mean([r['omega'] for r in results]),
        'mean_C': mean([r['C'] for r in results]),
        'mean_kappa': mean([r['kappa'] for r in results]),
    }
    with open(os.path.join(args.out_dir, 'summary.json'), 'w') as f:
        json.dump(summary, f, indent=2)
    plot_path = os.path.join(args.out_dir, 'density.png')
    density_plot(results, plot_path)
    docx_path = os.path.join(args.out_dir, 'report.docx')
    export_docx(summary, plot_path, docx_path)
    print(f"Generated report: {docx_path}")

if __name__ == '__main__':
    main()
