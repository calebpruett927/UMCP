"""
Analysis module for UMCP.

This module provides high-level analysis functions for:
- Time series analysis and regime classification
- Parameter tuning and what-if analysis
- Statistical analysis and reporting
- Visualization support
"""

import csv
import json
import os
from typing import List, Dict, Any, Optional, Tuple
from statistics import mean
try:
    import matplotlib.pyplot as plt
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from ..core import compute_invariants, classify_regime


def analyze_time_series(x_raw: List[float], a: float, b: float, 
                       **kwargs) -> Dict[str, Any]:
    """
    Perform comprehensive time series analysis.
    
    Parameters
    ----------
    x_raw : List[float]
        Raw time series data
    a : float
        Normalization intercept
    b : float
        Normalization scale
    **kwargs
        Additional parameters for compute_invariants
        
    Returns
    -------
    Dict[str, Any]
        Comprehensive analysis results
    """
    # Compute all invariants
    results = compute_invariants(x_raw, a, b, **kwargs)
    
    # Extract sequences for analysis
    omega = [r['omega'] for r in results]
    F = [r['F'] for r in results]
    S = [r['S'] for r in results]
    C = [r['C'] for r in results]
    kappa = [r['kappa'] for r in results]
    IC = [r['IC'] for r in results]
    regimes = [r['regime'] for r in results]
    
    # Count regimes
    regime_counts = {}
    for regime in regimes:
        regime_counts[regime] = regime_counts.get(regime, 0) + 1
    
    # Compute statistics
    stats = {
        'n_samples': len(results),
        'regime_counts': regime_counts,
        'mean_omega': mean(omega) if omega else 0.0,
        'mean_F': mean(F) if F else 0.0,
        'mean_S': mean(S) if S else 0.0,
        'mean_C': mean(C) if C else 0.0,
        'mean_kappa': mean(kappa) if kappa else 0.0,
        'mean_IC': mean(IC) if IC else 0.0,
        'max_omega': max(omega) if omega else 0.0,
        'max_C': max(C) if C else 0.0,
        'min_IC': min(IC) if IC else 0.0,
    }
    
    # Identify critical periods
    critical_periods = []
    for i, r in enumerate(results):
        if r['regime'] in ['Collapse', 'Critical']:
            critical_periods.append(i)
    
    return {
        'results': results,
        'statistics': stats,
        'critical_periods': critical_periods,
        'parameters': {'a': a, 'b': b, **kwargs}
    }


def parse_csv_channel(file_path: str, channel: str) -> List[float]:
    """
    Parse a specific channel from a CSV file.
    
    Parameters
    ----------
    file_path : str
        Path to CSV file
    channel : str
        Column name to extract
        
    Returns
    -------
    List[float]
        Extracted values
    """
    x_raw = []
    with open(file_path, 'r', newline='') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            try:
                value = float(row[channel])
                x_raw.append(value)
            except (KeyError, ValueError):
                continue
    return x_raw


def density_plot(results: List[Dict[str, Any]], out_path: str, 
                title: str = "Density of omega vs curvature (C)") -> None:
    """
    Create a density plot of omega vs curvature.
    
    Parameters
    ----------
    results : List[Dict[str, Any]]
        Analysis results from compute_invariants
    out_path : str
        Output file path
    title : str
        Plot title
    """
    if not HAS_MATPLOTLIB:
        raise ImportError("matplotlib is required for plotting")
    
    omegas = [r['omega'] for r in results]
    Cs = [r['C'] for r in results]
    
    plt.figure(figsize=(10, 6))
    plt.hexbin(omegas, Cs, gridsize=50, cmap='Blues', mincnt=1)
    plt.xlabel('omega (drift)')
    plt.ylabel('C (curvature)')
    plt.title(title)
    plt.colorbar(label='Counts')
    
    # Add regime gate lines
    plt.axvline(0.038, color='orange', linestyle='--', label='Watch gate')
    plt.axvline(0.30, color='red', linestyle='--', label='Collapse gate')
    plt.axhline(0.14, color='green', linestyle='--', label='C gate')
    
    plt.legend()
    plt.tight_layout()
    plt.savefig(out_path, dpi=150, bbox_inches='tight')
    plt.close()


def regime_timeline_plot(results: List[Dict[str, Any]], out_path: str) -> None:
    """
    Create a timeline plot showing regime evolution.
    
    Parameters
    ----------
    results : List[Dict[str, Any]]
        Analysis results
    out_path : str
        Output file path
    """
    if not HAS_MATPLOTLIB:
        raise ImportError("matplotlib is required for plotting")
    
    times = [r['t'] for r in results]
    regimes = [r['regime'] for r in results]
    
    # Map regimes to colors
    regime_colors = {
        'Stable': 'green',
        'Watch': 'orange', 
        'Collapse': 'red',
        'Critical': 'darkred'
    }
    
    colors = [regime_colors.get(r, 'gray') for r in regimes]
    
    plt.figure(figsize=(12, 4))
    plt.scatter(times, [1]*len(times), c=colors, alpha=0.7, s=20)
    plt.xlabel('Time')
    plt.ylabel('Regime')
    plt.title('Regime Classification Timeline')
    plt.yticks([])
    
    # Add legend
    for regime, color in regime_colors.items():
        if regime in regimes:
            plt.scatter([], [], c=color, label=regime)
    plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left')
    
    plt.tight_layout()
    plt.savefig(out_path, dpi=150, bbox_inches='tight')
    plt.close()


def export_report(analysis: Dict[str, Any], plot_paths: List[str], 
                 out_docx: str) -> None:
    """
    Export analysis report to Word document.
    
    Parameters
    ----------
    analysis : Dict[str, Any]
        Analysis results from analyze_time_series
    plot_paths : List[str]
        Paths to plot images to include
    out_docx : str
        Output document path
    """
    if not HAS_DOCX:
        raise ImportError("python-docx is required for document export")
    
    doc = Document()
    doc.add_heading('UMCP Analysis Report', level=1)
    
    stats = analysis['statistics']
    params = analysis['parameters']
    
    # Parameters section
    doc.add_heading('Parameters', level=2)
    doc.add_paragraph(f"Normalization: a={params['a']:.6f}, b={params['b']:.6f}")
    for key, value in params.items():
        if key not in ['a', 'b']:
            doc.add_paragraph(f"{key}: {value}")
    
    # Statistics section
    doc.add_heading('Statistics', level=2)
    doc.add_paragraph(f"Total samples: {stats['n_samples']}")
    doc.add_paragraph(f"Mean omega: {stats['mean_omega']:.4f}")
    doc.add_paragraph(f"Mean curvature C: {stats['mean_C']:.4f}")
    doc.add_paragraph(f"Mean kappa: {stats['mean_kappa']:.4f}")
    doc.add_paragraph(f"Mean IC: {stats['mean_IC']:.4f}")
    
    # Regime counts
    doc.add_heading('Regime Distribution', level=2)
    for regime, count in stats['regime_counts'].items():
        percentage = (count / stats['n_samples']) * 100
        doc.add_paragraph(f"{regime}: {count} ({percentage:.1f}%)", style='List Bullet')
    
    # Critical periods
    if analysis['critical_periods']:
        doc.add_heading('Critical Periods', level=2)
        doc.add_paragraph(f"Critical events detected at time steps: {analysis['critical_periods']}")
    
    # Add plots
    for i, plot_path in enumerate(plot_paths):
        if os.path.exists(plot_path):
            if i == 0:
                doc.add_heading('Visualizations', level=2)
            try:
                # Scale image to fit page width
                doc.add_picture(plot_path, width=doc.sections[0].page_width * 0.8)
                doc.add_paragraph("")  # Add space after image
            except Exception as e:
                doc.add_paragraph(f"Could not include plot: {plot_path} (Error: {e})")
    
    doc.save(out_docx)


def parameter_sweep(x_raw: List[float], a_values: List[float], b_values: List[float],
                   **kwargs) -> List[Dict[str, Any]]:
    """
    Perform parameter sweep analysis across different (a,b) values.
    
    Parameters
    ----------
    x_raw : List[float]
        Raw time series data
    a_values : List[float]
        List of 'a' values to test
    b_values : List[float]
        List of 'b' values to test
    **kwargs
        Additional parameters for analysis
        
    Returns
    -------
    List[Dict[str, Any]]
        Results for each parameter combination
    """
    sweep_results = []
    
    for a in a_values:
        for b in b_values:
            try:
                analysis = analyze_time_series(x_raw, a, b, **kwargs)
                sweep_results.append({
                    'a': a,
                    'b': b,
                    'statistics': analysis['statistics'],
                    'n_critical': len(analysis['critical_periods'])
                })
            except Exception as e:
                sweep_results.append({
                    'a': a,
                    'b': b,
                    'error': str(e)
                })
    
    return sweep_results


def playground_analysis(csv_path: str, channel: str, a: float, b: float,
                       out_dir: str = 'playground_out', **kwargs) -> str:
    """
    Run complete playground analysis on CSV data.
    
    Parameters
    ----------
    csv_path : str
        Path to CSV file
    channel : str
        Column name to analyze
    a : float
        Normalization intercept
    b : float
        Normalization scale
    out_dir : str
        Output directory
    **kwargs
        Additional analysis parameters
        
    Returns
    -------
    str
        Path to generated report
    """
    # Create output directory
    os.makedirs(out_dir, exist_ok=True)
    
    # Parse data
    x_raw = parse_csv_channel(csv_path, channel)
    if not x_raw:
        raise ValueError(f"No data found in column '{channel}' of {csv_path}")
    
    # Run analysis
    analysis = analyze_time_series(x_raw, a, b, **kwargs)
    
    # Save summary JSON
    summary_path = os.path.join(out_dir, 'summary.json')
    with open(summary_path, 'w') as f:
        # Make a JSON-serializable copy
        json_analysis = {
            'statistics': analysis['statistics'],
            'parameters': analysis['parameters'],
            'n_critical_periods': len(analysis['critical_periods']),
            'critical_periods': analysis['critical_periods']
        }
        json.dump(json_analysis, f, indent=2)
    
    # Generate plots
    plot_paths = []
    
    if HAS_MATPLOTLIB:
        try:
            # Density plot
            density_path = os.path.join(out_dir, 'density.png')
            density_plot(analysis['results'], density_path)
            plot_paths.append(density_path)
            
            # Timeline plot
            timeline_path = os.path.join(out_dir, 'timeline.png')
            regime_timeline_plot(analysis['results'], timeline_path)
            plot_paths.append(timeline_path)
        except Exception as e:
            print(f"Warning: Could not generate plots: {e}")
    
    # Export report
    report_path = os.path.join(out_dir, 'report.docx')
    try:
        export_report(analysis, plot_paths, report_path)
        print(f"Generated report: {report_path}")
        return report_path
    except Exception as e:
        print(f"Warning: Could not generate Word report: {e}")
        return summary_path


# Export main functions
__all__ = [
    "analyze_time_series", "parse_csv_channel", "density_plot", 
    "regime_timeline_plot", "export_report", "parameter_sweep",
    "playground_analysis"
]